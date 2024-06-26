import docx

def say_hello():
    print("Hello world!!!!!!!!!")

say_hello()

def open_docx():
    # f = open("files/test.docx")
    # lines = f.readlines()
    # print(lines)
    # f.close()

    document = docx.Document("files/test.docx")
    # # for paragraph in word.paragraphs:
    for paragraph in document.paragraphs:
        print(paragraph.text)

open_docx()